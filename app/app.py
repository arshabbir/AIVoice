import json
import logging
import os
import re
import shutil
import sqlite3
import subprocess
import time
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable

import numpy as np
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from fastapi import FastAPI, File, HTTPException, Request, UploadFile
from fastapi.responses import HTMLResponse, JSONResponse, RedirectResponse
from fastapi.templating import Jinja2Templates
from livekit import api
from pydantic import BaseModel
from pypdf import PdfReader
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, PointStruct, VectorParams
from sentence_transformers import SentenceTransformer


def setup_logging() -> logging.Logger:
    level_name = os.getenv("APP_LOG_LEVEL", "INFO").upper()
    level = getattr(logging, level_name, logging.INFO)

    logger = logging.getLogger("bpo-app")
    logger.setLevel(level)
    logger.handlers.clear()

    handler = logging.StreamHandler()
    handler.setLevel(level)
    handler.setFormatter(
        logging.Formatter("%(asctime)s %(levelname)s [%(name)s] %(message)s")
    )
    logger.addHandler(handler)
    logger.propagate = False
    return logger


logger = setup_logging()

APP_PUBLIC_URL = os.getenv("APP_PUBLIC_URL", "http://localhost:8080")
LIVEKIT_PUBLIC_URL = os.getenv("LIVEKIT_PUBLIC_URL", "ws://localhost:7880")
LIVEKIT_API_KEY = os.getenv("LIVEKIT_API_KEY", "devkey")
LIVEKIT_API_SECRET = os.getenv("LIVEKIT_API_SECRET", "devsecret")

UPLOAD_DIR = Path(os.getenv("UPLOAD_DIR", "/data/uploads"))
DB_PATH = Path(os.getenv("APP_DB_PATH", "/data/app.db"))

QDRANT_URL = os.getenv("QDRANT_URL", "http://qdrant:6333")
QDRANT_COLLECTION = os.getenv("QDRANT_COLLECTION", "knowledge_base")

EMBED_MODEL = os.getenv("EMBED_MODEL", "sentence-transformers/all-MiniLM-L6-v2")
EMBED_BATCH_SIZE = int(os.getenv("EMBED_BATCH_SIZE", "64"))

SEARCH_TOP_K = int(os.getenv("SEARCH_TOP_K", "4"))
RAG_CONFIDENCE_THRESHOLD = float(os.getenv("RAG_CONFIDENCE_THRESHOLD", "0.23"))
ANSWER_MODE = os.getenv("ANSWER_MODE", "hybrid")

CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", "700"))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", "120"))

QDRANT_UPSERT_BATCH_SIZE = int(os.getenv("QDRANT_UPSERT_BATCH_SIZE", "64"))
QDRANT_MAX_REQUEST_BYTES = int(
    os.getenv("QDRANT_MAX_REQUEST_BYTES", str(28 * 1024 * 1024))
)

UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
DB_PATH.parent.mkdir(parents=True, exist_ok=True)

app = FastAPI(title="BPO Voice Agent App")
templates = Jinja2Templates(directory="templates")


def build_qdrant_client() -> QdrantClient:
    return QdrantClient(url=QDRANT_URL)


def wait_for_qdrant(attempts: int = 60, delay: float = 2.0) -> QdrantClient:
    last_exc = None
    for _ in range(attempts):
        try:
            client = build_qdrant_client()
            client.get_collections()
            return client
        except Exception as exc:
            last_exc = exc
            time.sleep(delay)
    raise RuntimeError(f"Qdrant is not reachable at {QDRANT_URL}: {last_exc}")


qdrant = wait_for_qdrant()
embedder = SentenceTransformer(EMBED_MODEL)
VECTOR_DIM = int(embedder.get_sentence_embedding_dimension())


class SearchRequest(BaseModel):
    query: str
    top_k: int = SEARCH_TOP_K


class HandoffRequest(BaseModel):
    room_name: str
    reason: str


class TokenRequest(BaseModel):
    room_name: str | None = None
    participant_identity: str | None = None
    participant_name: str | None = None
    participant_metadata: str | None = None


def utcnow() -> str:
    return datetime.now(timezone.utc).isoformat()


def db() -> sqlite3.Connection:
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def ensure_column(conn: sqlite3.Connection, table: str, column: str, ddl: str) -> None:
    cols = {row["name"] for row in conn.execute(f"PRAGMA table_info({table})").fetchall()}
    if column not in cols:
        conn.execute(f"ALTER TABLE {table} ADD COLUMN {ddl}")


def count_documents(conn: sqlite3.Connection | None = None) -> int:
    owns_conn = conn is None
    if owns_conn:
        conn = db()
    try:
        row = conn.execute(
            "SELECT COUNT(*) AS cnt FROM documents WHERE status='ready'"
        ).fetchone()
        return int(row["cnt"] if row else 0)
    finally:
        if owns_conn and conn is not None:
            conn.close()


def init_db() -> None:
    conn = db()
    cur = conn.cursor()

    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS documents (
            id TEXT PRIMARY KEY,
            filename TEXT NOT NULL,
            original_path TEXT NOT NULL,
            uploaded_at TEXT NOT NULL,
            chunk_count INTEGER NOT NULL DEFAULT 0
        )
        """
    )
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS handoffs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            room_name TEXT NOT NULL,
            reason TEXT NOT NULL,
            created_at TEXT NOT NULL,
            status TEXT NOT NULL DEFAULT 'open'
        )
        """
    )

    ensure_column(conn, "documents", "status", "status TEXT NOT NULL DEFAULT 'ready'")
    ensure_column(conn, "documents", "error_message", "error_message TEXT")

    conn.commit()
    conn.close()


def init_qdrant() -> None:
    try:
        qdrant.get_collection(QDRANT_COLLECTION)
    except Exception:
        qdrant.create_collection(
            collection_name=QDRANT_COLLECTION,
            vectors_config=VectorParams(size=VECTOR_DIM, distance=Distance.COSINE),
        )


def qdrant_search_compat(
    *,
    client: QdrantClient,
    collection_name: str,
    query_vector: list[float],
    limit: int,
):
    if hasattr(client, "search"):
        return client.search(
            collection_name=collection_name,
            query_vector=query_vector,
            limit=limit,
            with_payload=True,
        )

    if hasattr(client, "query_points"):
        result = client.query_points(
            collection_name=collection_name,
            query=query_vector,
            limit=limit,
            with_payload=True,
        )
        if hasattr(result, "points"):
            return result.points
        if isinstance(result, dict) and "points" in result:
            return result["points"]
        return result

    raise RuntimeError("Unsupported qdrant-client version: no search/query_points API found")


def hit_payload(hit: Any) -> dict[str, Any]:
    if hasattr(hit, "payload"):
        return getattr(hit, "payload") or {}
    if isinstance(hit, dict):
        return hit.get("payload", {}) or {}
    return {}


def hit_score(hit: Any) -> float:
    if hasattr(hit, "score"):
        return float(getattr(hit, "score") or 0.0)
    if isinstance(hit, dict):
        return float(hit.get("score", 0.0) or 0.0)
    return 0.0


def point_to_primitive(point: PointStruct) -> dict[str, Any]:
    if hasattr(point, "model_dump"):
        return point.model_dump()
    if hasattr(point, "dict"):
        return point.dict()
    return {
        "id": getattr(point, "id", None),
        "vector": getattr(point, "vector", None),
        "payload": getattr(point, "payload", None),
    }


def estimate_point_payload_bytes(point: PointStruct) -> int:
    primitive = point_to_primitive(point)
    return len(json.dumps(primitive, separators=(",", ":"), ensure_ascii=False).encode("utf-8"))


def split_points_for_qdrant(
    points: list[PointStruct],
    *,
    max_points: int,
    max_bytes: int,
) -> list[list[PointStruct]]:
    batches: list[list[PointStruct]] = []
    current_batch: list[PointStruct] = []
    current_bytes = 0

    for point in points:
        point_bytes = estimate_point_payload_bytes(point)

        if point_bytes > max_bytes:
            raise HTTPException(
                status_code=400,
                detail=(
                    "A single indexed chunk is too large for Qdrant request limits. "
                    "Reduce chunk size or trim document content."
                ),
            )

        should_flush = False
        if current_batch and len(current_batch) >= max_points:
            should_flush = True
        if current_batch and (current_bytes + point_bytes) > max_bytes:
            should_flush = True

        if should_flush:
            batches.append(current_batch)
            current_batch = []
            current_bytes = 0

        current_batch.append(point)
        current_bytes += point_bytes

    if current_batch:
        batches.append(current_batch)

    return batches


def upsert_points_in_batches(
    *,
    points: list[PointStruct],
    document_id: str,
    trace_id: str,
) -> None:
    if not points:
        return

    batches = split_points_for_qdrant(
        points,
        max_points=QDRANT_UPSERT_BATCH_SIZE,
        max_bytes=QDRANT_MAX_REQUEST_BYTES,
    )

    for idx, batch in enumerate(batches, start=1):
        approx_bytes = sum(estimate_point_payload_bytes(p) for p in batch)
        logger.info(
            "qdrant_upsert_batch trace_id=%s document_id=%s batch=%s/%s points=%s approx_payload_bytes=%s",
            trace_id,
            document_id,
            idx,
            len(batches),
            len(batch),
            approx_bytes,
        )
        qdrant.upsert(
            collection_name=QDRANT_COLLECTION,
            wait=True,
            points=batch,
        )


init_db()
init_qdrant()


@app.middleware("http")
async def request_logging(request: Request, call_next):
    started = time.perf_counter()
    trace_id = request.headers.get("X-Trace-Id") or uuid.uuid4().hex[:10]
    request.state.trace_id = trace_id
    try:
        response = await call_next(request)
    except Exception:
        logger.exception(
            "request_failed trace_id=%s method=%s path=%s",
            trace_id,
            request.method,
            request.url.path,
        )
        raise
    elapsed_ms = round((time.perf_counter() - started) * 1000, 1)
    response.headers["X-Trace-Id"] = trace_id
    logger.info(
        "request_completed trace_id=%s method=%s path=%s status=%s elapsed_ms=%s",
        trace_id,
        request.method,
        request.url.path,
        response.status_code,
        elapsed_ms,
    )
    return response


@app.get("/")
def root() -> RedirectResponse:
    return RedirectResponse(url="/admin")


@app.get("/api/health")
def api_health() -> dict[str, Any]:
    return {
        "status": "ok",
        "answer_mode": ANSWER_MODE,
        "documents": count_documents(),
        "qdrant_url": QDRANT_URL,
        "chunk_size": CHUNK_SIZE,
        "chunk_overlap": CHUNK_OVERLAP,
        "upsert_batch_size": QDRANT_UPSERT_BATCH_SIZE,
    }


@app.get("/admin", response_class=HTMLResponse)
def admin_page(request: Request):
    return templates.TemplateResponse(
        "admin.html",
        {
            "request": request,
            "app_public_url": APP_PUBLIC_URL,
            "answer_mode": ANSWER_MODE,
        },
    )


@app.get("/test", response_class=HTMLResponse)
def test_page(request: Request, room: str | None = None, identity: str | None = None):
    return templates.TemplateResponse(
        "test.html",
        {
            "request": request,
            "default_room": room or "",
            "default_identity": identity or "",
            "app_public_url": APP_PUBLIC_URL,
            "answer_mode": ANSWER_MODE,
        },
    )


@app.post("/api/token", status_code=201)
def issue_token(req: TokenRequest, request: Request):
    room_name = req.room_name or f"room-{uuid.uuid4().hex[:8]}"
    participant_identity = req.participant_identity or f"user-{uuid.uuid4().hex[:8]}"
    participant_name = req.participant_name or participant_identity
    participant_metadata = req.participant_metadata or ""

    token = (
        api.AccessToken(LIVEKIT_API_KEY, LIVEKIT_API_SECRET)
        .with_identity(participant_identity)
        .with_name(participant_name)
        .with_metadata(participant_metadata)
        .with_grants(
            api.VideoGrants(
                room_join=True,
                room=room_name,
                can_publish=True,
                can_publish_data=True,
                can_subscribe=True,
            )
        )
        .to_jwt()
    )

    logger.info(
        "token_issued trace_id=%s room=%s identity=%s answer_mode=%s",
        getattr(request.state, "trace_id", "n/a"),
        room_name,
        participant_identity,
        ANSWER_MODE,
    )

    return JSONResponse(
        status_code=201,
        content={
            "server_url": LIVEKIT_PUBLIC_URL,
            "participant_token": token,
            "room_name": room_name,
            "participant_identity": participant_identity,
        },
    )


def normalize_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def chunk_text(
    text: str,
    chunk_size: int = CHUNK_SIZE,
    overlap: int = CHUNK_OVERLAP,
) -> list[str]:
    text = normalize_text(text)
    if not text:
        return []

    chunks: list[str] = []
    start = 0
    text_len = len(text)

    while start < text_len:
        end = min(text_len, start + chunk_size)
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        if end >= text_len:
            break

        start = max(0, end - overlap)

    return chunks


def read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def read_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    parts: list[str] = []

    for paragraph in doc.paragraphs:
        if paragraph.text:
            parts.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


def read_doc(path: Path) -> str:
    antiword = shutil.which("antiword")
    if not antiword:
        raise HTTPException(
            status_code=400,
            detail=(
                "Legacy .doc files require 'antiword' in the app container. "
                "Install antiword or convert the file to .docx."
            ),
        )

    result = subprocess.run(
        [antiword, str(path)],
        capture_output=True,
        text=True,
        check=False,
    )

    if result.returncode != 0:
        stderr = (result.stderr or "").strip()
        raise HTTPException(
            status_code=400,
            detail=f"Failed to parse .doc file with antiword: {stderr or 'unknown error'}",
        )

    return result.stdout or ""


def read_textlike(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def read_html(path: Path) -> str:
    html = path.read_text(encoding="utf-8", errors="ignore")
    soup = BeautifulSoup(html, "html.parser")
    return soup.get_text(" ")


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return read_pdf(path)
    if suffix == ".docx":
        return read_docx(path)
    if suffix == ".doc":
        return read_doc(path)
    if suffix in {".txt", ".md", ".csv", ".json", ".log"}:
        return read_textlike(path)
    if suffix in {".html", ".htm"}:
        return read_html(path)

    raise HTTPException(status_code=400, detail=f"Unsupported file type: {suffix}")


def embed_texts(texts: Iterable[str]) -> np.ndarray:
    items = list(texts)
    all_vectors: list[np.ndarray] = []

    for i in range(0, len(items), EMBED_BATCH_SIZE):
        batch = items[i : i + EMBED_BATCH_SIZE]
        vectors = embedder.encode(batch, normalize_embeddings=True)
        all_vectors.append(np.asarray(vectors, dtype=np.float32))

    if not all_vectors:
        return np.empty((0, VECTOR_DIM), dtype=np.float32)

    return np.vstack(all_vectors)


def update_document_status(
    document_id: str,
    *,
    status: str,
    chunk_count: int | None = None,
    error_message: str | None = None,
) -> None:
    conn = db()
    if chunk_count is None:
        conn.execute(
            "UPDATE documents SET status=?, error_message=? WHERE id=?",
            (status, error_message, document_id),
        )
    else:
        conn.execute(
            "UPDATE documents SET status=?, chunk_count=?, error_message=? WHERE id=?",
            (status, chunk_count, error_message, document_id),
        )
    conn.commit()
    conn.close()


@app.post("/api/upload")
async def upload_document(request: Request, file: UploadFile = File(...)):
    if not file.filename:
        raise HTTPException(status_code=400, detail="filename is required")

    started = time.perf_counter()
    trace_id = getattr(request.state, "trace_id", "n/a")
    doc_id = uuid.uuid4().hex
    original_name = Path(file.filename).name
    stored_path = UPLOAD_DIR / f"{doc_id}{Path(file.filename).suffix.lower()}"

    payload = await file.read()
    stored_path.write_bytes(payload)

    conn = db()
    conn.execute(
        """
        INSERT INTO documents (id, filename, original_path, uploaded_at, chunk_count, status, error_message)
        VALUES (?, ?, ?, ?, ?, ?, ?)
        """,
        (doc_id, original_name, str(stored_path), utcnow(), 0, "processing", None),
    )
    conn.commit()
    conn.close()

    try:
        text = extract_text(stored_path)
        chunks = chunk_text(text)
        if not chunks:
            raise HTTPException(status_code=400, detail="No extractable text found in document")

        logger.info(
            "document_chunking_completed trace_id=%s document_id=%s filename=%s bytes=%s chunks=%s chunk_size=%s overlap=%s",
            trace_id,
            doc_id,
            original_name,
            len(payload),
            len(chunks),
            CHUNK_SIZE,
            CHUNK_OVERLAP,
        )

        vectors = embed_texts(chunks)

        doc_uuid = uuid.uuid5(uuid.NAMESPACE_URL, doc_id)
        points: list[PointStruct] = []

        for idx, (chunk, vector) in enumerate(zip(chunks, vectors)):
            point_id = str(uuid.uuid5(doc_uuid, f"chunk-{idx}"))
            points.append(
                PointStruct(
                    id=point_id,
                    vector=vector.tolist(),
                    payload={
                        "document_id": doc_id,
                        "filename": original_name,
                        "chunk_index": idx,
                        "text": chunk,
                    },
                )
            )

        upsert_points_in_batches(
            points=points,
            document_id=doc_id,
            trace_id=trace_id,
        )

        update_document_status(
            doc_id,
            status="ready",
            chunk_count=len(chunks),
            error_message=None,
        )

        elapsed_ms = round((time.perf_counter() - started) * 1000, 1)
        logger.info(
            "document_indexed trace_id=%s document_id=%s filename=%s bytes=%s chunks=%s elapsed_ms=%s",
            trace_id,
            doc_id,
            original_name,
            len(payload),
            len(chunks),
            elapsed_ms,
        )

        return {
            "document_id": doc_id,
            "filename": original_name,
            "chunks": len(chunks),
            "elapsed_ms": elapsed_ms,
            "status": "ready",
        }

    except HTTPException as exc:
        update_document_status(
            doc_id,
            status="failed",
            chunk_count=0,
            error_message=str(exc.detail),
        )
        logger.warning(
            "document_index_failed trace_id=%s document_id=%s filename=%s detail=%s",
            trace_id,
            doc_id,
            original_name,
            exc.detail,
        )
        raise
    except Exception as exc:
        update_document_status(
            doc_id,
            status="failed",
            chunk_count=0,
            error_message=str(exc),
        )
        logger.exception(
            "document_index_failed trace_id=%s document_id=%s filename=%s error=%s",
            trace_id,
            doc_id,
            original_name,
            exc,
        )
        raise HTTPException(status_code=500, detail=f"Failed to index document: {exc}") from exc


@app.get("/api/documents")
def list_documents():
    conn = db()
    rows = conn.execute(
        """
        SELECT id, filename, uploaded_at, chunk_count, status, error_message
        FROM documents
        ORDER BY uploaded_at DESC
        """
    ).fetchall()
    conn.close()
    return [dict(row) for row in rows]


@app.post("/api/search")
def search(req: SearchRequest, request: Request):
    query = normalize_text(req.query)
    if not query:
        raise HTTPException(status_code=400, detail="query is required")

    started = time.perf_counter()
    trace_id = getattr(request.state, "trace_id", "n/a")
    total_documents = count_documents()

    if total_documents == 0:
        elapsed_ms = round((time.perf_counter() - started) * 1000, 1)
        logger.info(
            "search_completed trace_id=%s query=%r snippets=0 confidence=0.000 total_documents=0 elapsed_ms=%s",
            trace_id,
            query,
            elapsed_ms,
        )
        return {
            "query": query,
            "confidence": 0.0,
            "snippets": [],
            "total_documents": 0,
            "elapsed_ms": elapsed_ms,
            "backend": "qdrant",
            "threshold": RAG_CONFIDENCE_THRESHOLD,
        }

    try:
        vector = embed_texts([query])[0].tolist()
        hits = qdrant_search_compat(
            client=qdrant,
            collection_name=QDRANT_COLLECTION,
            query_vector=vector,
            limit=max(1, req.top_k),
        )
    except Exception as exc:
        elapsed_ms = round((time.perf_counter() - started) * 1000, 1)
        logger.exception(
            "search_failed trace_id=%s query=%r total_documents=%s elapsed_ms=%s error=%s",
            trace_id,
            query,
            total_documents,
            elapsed_ms,
            exc,
        )
        return {
            "query": query,
            "confidence": 0.0,
            "snippets": [],
            "total_documents": total_documents,
            "elapsed_ms": elapsed_ms,
            "backend": "fallback",
            "error": "qdrant_search_failed",
        }

    snippets = []
    for hit in hits or []:
        payload = hit_payload(hit)
        snippets.append(
            {
                "document_id": payload.get("document_id"),
                "filename": payload.get("filename"),
                "chunk_index": payload.get("chunk_index"),
                "text": payload.get("text", ""),
                "score": hit_score(hit),
            }
        )

    confidence = float(snippets[0]["score"]) if snippets else 0.0
    confidence = max(0.0, min(1.0, confidence))
    elapsed_ms = round((time.perf_counter() - started) * 1000, 1)

    logger.info(
        "search_completed trace_id=%s query=%r snippets=%s confidence=%.3f total_documents=%s elapsed_ms=%s",
        trace_id,
        query,
        len(snippets),
        confidence,
        total_documents,
        elapsed_ms,
    )

    return {
        "query": query,
        "confidence": confidence,
        "snippets": snippets,
        "total_documents": total_documents,
        "elapsed_ms": elapsed_ms,
        "backend": "qdrant",
        "threshold": RAG_CONFIDENCE_THRESHOLD,
    }


@app.post("/api/handoff")
def create_handoff(req: HandoffRequest, request: Request):
    conn = db()
    cur = conn.execute(
        "INSERT INTO handoffs (room_name, reason, created_at, status) VALUES (?, ?, ?, 'open')",
        (req.room_name, req.reason, utcnow()),
    )
    conn.commit()
    handoff_id = cur.lastrowid
    conn.close()

    logger.info(
        "handoff_created trace_id=%s handoff_id=%s room=%s reason=%r",
        getattr(request.state, "trace_id", "n/a"),
        handoff_id,
        req.room_name,
        req.reason,
    )
    return {"id": handoff_id, "status": "open"}


@app.get("/api/handoffs")
def list_handoffs():
    conn = db()
    rows = conn.execute(
        "SELECT id, room_name, reason, created_at, status FROM handoffs ORDER BY id DESC LIMIT 100"
    ).fetchall()
    conn.close()
    return [dict(row) for row in rows]


@app.post("/api/handoffs/{handoff_id}/close")
def close_handoff(handoff_id: int, request: Request):
    conn = db()
    conn.execute("UPDATE handoffs SET status='closed' WHERE id=?", (handoff_id,))
    conn.commit()
    conn.close()

    logger.info(
        "handoff_closed trace_id=%s handoff_id=%s",
        getattr(request.state, "trace_id", "n/a"),
        handoff_id,
    )
    return {"id": handoff_id, "status": "closed"}